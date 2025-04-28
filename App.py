from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import random
import spacy
from spacy.lang.fr.stop_words import STOP_WORDS
import streamlit as st
import re
import string
from io import BytesIO
import pandas as pd
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import pdfminer.high_level
import docx
from openai import AzureOpenAI
import os

# Configuration de l'application
st.set_page_config(page_title="Générateur de Cas de test à partir du CDC", layout="wide", page_icon="📑")

# ----------------------------
# FONCTIONS UTILITAIRES
# ----------------------------

@st.cache_resource
def load_nlp_model():
    """Charge le modèle spaCy pour le traitement NLP"""
    try:
        # Essaye de charger le modèle normalement
        nlp = spacy.load("fr_core_news_md")
        st.success("Modèle NLP chargé avec succès !")
        return nlp
    except OSError:
        try:
            # Si le modèle n'est pas trouvé, propose l'installation
            st.error("Modèle français non trouvé. Installation en cours...")
            import os
            os.system("python -m spacy download fr_core_news_md")
            nlp = spacy.load("fr_core_news_md")
            return nlp
        except Exception as e:
            st.error(f"Échec du chargement : {str(e)}")
            return None
    except Exception as e:
        st.error(f"Erreur inattendue : {str(e)}")
        return None

def setup_azure_openai():
    """Version ultra-robuste avec fallback"""
    try:
        # Méthode 1 : Lecture depuis secrets.toml
        if hasattr(st, 'secrets') and 'azure_openai' in st.secrets:
            config = st.secrets["azure_openai"]
        # Méthode 2 : Fallback pour développement local
        else:
            from dotenv import load_dotenv
            load_dotenv()
            config = {
                "AZURE_OPENAI_KEY": os.getenv("AZURE_OPENAI_KEY"),
                "AZURE_OPENAI_ENDPOINT": os.getenv("AZURE_OPENAI_ENDPOINT"),
                "DEPLOYMENT_NAME": os.getenv("DEPLOYMENT_NAME", "gpt-4o"),
                "API_VERSION": os.getenv("API_VERSION", "2024-02-15-preview")
            }

        client = AzureOpenAI(
            api_key=config["AZURE_OPENAI_KEY"],
            api_version=config["API_VERSION"],  # Utilisation de la version depuis les secrets
            azure_endpoint=config["AZURE_OPENAI_ENDPOINT"]
        )
        
        # Test de connexion immédiat
        test = client.chat.completions.create(
            model=config["DEPLOYMENT_NAME"],
            messages=[{"role": "user", "content": "Test"}],
            max_tokens=5
        )
        return client
        
    except Exception as e:
        st.error(f"""
        ❌ ERREUR CRITIQUE Azure OpenAI :
        {str(e)}
        
        🔍 Vérifiez que :
        1. Le fichier .streamlit/secrets.toml existe
        2. Les clés sont correctes
        3. Le endpoint est accessible
        4. La version de l'API est valide
        """)
        return None
    
def generate_with_azure_openai(prompt, client, model="gpt-4", temperature=0.7, max_tokens=1000):
    """Génère du texte avec Azure OpenAI"""
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=temperature,
            max_tokens=max_tokens
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Erreur avec Azure OpenAI: {str(e)}")
        return None

def extract_business_rules(text, nlp_model, use_ai=False):
    """
    Extrait les règles métier du texte avec option pour utiliser Azure OpenAI
    """
    if not use_ai:
        # Méthode originale avec regex et NLP
        patterns = [
            r"(Si|Lorsqu'|Quand|Dès que|En cas de).*?(alors|doit|devra|est tenu de|nécessite|implique|entraîne|peut).*?\.",
            r"(Tout utilisateur|L'[a-zA-Z]+|Un client|Le système|Une demande).*?(doit|est tenu de|devra|ne peut pas|ne doit pas|est interdit de).*?\.",
            r"(Le non-respect|Toute infraction|Une violation).*?(entraîne|provoque|peut entraîner|résulte en|sera soumis à).*?\.",
            r"(L'utilisateur|Le client|Le prestataire|L'agent|Le système).*?(est autorisé à|peut|a le droit de).*?\."
        ]
        
        rules = set()
        
        # Extraction par motifs regex
        for pattern in patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                rules.add(clean_rule(match.group()))
        
        # Extraction NLP si le modèle est disponible
        if nlp_model:
            doc = nlp_model(text)
            for sent in doc.sents:
                if any(keyword in sent.text.lower() for keyword in ["si ", "alors", "doit", "est tenu de", "ne peut pas", "entraîne", "provoque",
                "peut entraîner", "doit être", "est obligatoire", "a le droit de", "est autorisé à"]):
                    if len(sent.text.split()) > 5:
                        rules.add(clean_rule(sent.text))
        
        return sorted(rules, key=lambda x: len(x), reverse=True)
    else:
        # Méthode avec Azure OpenAI
        client = setup_azure_openai()
        if not client:
            return []
            
        prompt = f"""
        Extrait les règles de gestion métier à partir du texte suivant en suivant ces consignes:
        1. Identifie toutes les règles fonctionnelles
        2. Formule-les de manière claire et concise
        3. Chaque règle doit être autonome et complète
        4. Conserve la formulation originale autant que possible
        5. Retourne uniquement les règles, une par ligne
        
        Texte:
        {text[:10000]}  # Limite pour éviter les tokens excessifs
        """
        
        result = generate_with_azure_openai(prompt, client)
        if result:
            return [clean_rule(rule) for rule in result.split('\n') if rule.strip()]
        return []

def clean_rule(rule_text):
    """Nettoie et formate une règle de gestion"""
    rule_text = re.sub(r"\s+", " ", rule_text).strip()
    if not rule_text.endswith('.'):
        rule_text += '.'
    return rule_text

def extract_text(uploaded_file):
    """Extrait le texte depuis PDF ou DOCX"""
    try:
        file_bytes = uploaded_file.getvalue()
        
        if uploaded_file.type == "application/pdf":
            with BytesIO(file_bytes) as f:
                text = pdfminer.high_level.extract_text(f)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with BytesIO(file_bytes) as f:
                doc = docx.Document(f)
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            raise ValueError("Format non supporté")
            
        return text if text and text.strip() else None
        
    except Exception as e:
        st.error(f"Erreur d'extraction : {str(e)}")
        return None

def create_rules_document(rules):
    """Crée un document Word des règles"""
    doc = Document()
    doc.add_heading('Règles de Gestion Identifiées', level=1)
    
    for i, rule in enumerate(rules, 1):
        doc.add_paragraph(f"{i}. {rule}", style='ListBullet')
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def clean_text(text, nlp_model, min_word_length=3):
    """
    Nettoyage approfondi du texte avec :
    - Suppression des stopwords
    - Lemmatisation
    - Filtrage par catégorie grammaticale
    - Suppression des mots trop courts
    """
    if not text or not nlp_model:
        return ""
    
    # Nettoyage de base
    text = text.lower()
    text = re.sub(r"[^\w\sàâäéèêëîïôöùûüç]", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    
    # Traitement NLP
    doc = nlp_model(text)
    cleaned_tokens = []
    
    for token in doc:
        if (token.is_stop or 
            token.is_punct or 
            len(token.text) < min_word_length or
            token.pos_ in ["DET", "ADP", "CCONJ", "PRON", "PART"]):
            continue
            
        lemma = token.lemma_.strip()
        if lemma:
            cleaned_tokens.append(lemma)
    
    return " ".join(cleaned_tokens)

def calculate_frequencies(text):
    """Calcule les fréquences des mots"""
    words = [word for word in text.split() if len(word) > 2]
    return pd.Series(words).value_counts()

def generate_wordcloud(freq_dict, width=800, height=400, background_color="white", colormap="viridis"):
    """Génère un nuage de mots"""
    fig, ax = plt.subplots(figsize=(10, 5))
    wc = WordCloud(
        width=width,
        height=height,
        background_color=background_color,
        colormap=colormap,
        max_words=100
    ).generate_from_frequencies(freq_dict)
    
    ax.imshow(wc, interpolation="bilinear")
    ax.axis("off")
    return fig

def extract_pdc_from_text(text):
    """Extrait les exigences PDC d'un texte"""
    patterns = [
        r"(Vérifier|S['']assurer|Contrôler|Vérification|Point de contrôle)\b.*?[\.;]",
        r"(Le système doit|Il faut|Il est nécessaire de).*?(vérifier|contrôler|s'assurer)"
    ]
    pdc_list = set()
    for pattern in patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            pdc = match.group().strip()
            if len(pdc.split()) > 3:
                if not pdc.endswith('.'):
                    pdc += '.'
                pdc_list.add(pdc)
    return sorted(pdc_list, key=lambda x: len(x), reverse=True)

def generate_pdc_from_rule(rule, use_ai=False):
    """Génère un PDC à partir d'une règle de gestion"""
    if not use_ai or 'nlp' not in st.session_state:
        if 'nlp' not in st.session_state:
            st.error("Modèle NLP non chargé")
            return f"Vérifier que {rule}"
        
        doc = st.session_state.nlp(rule)
        verbs = [token.text for token in doc if token.pos_ == "VERB"]
        action = verbs[0] if verbs else "vérifier"
        return f"{action.capitalize()} que {rule}"
    else:
        client = setup_azure_openai()
        if not client:
            return f"Vérifier que {rule}"
            
        prompt = f"""
        Transforme cette règle de gestion en Point de Contrôle (PDC) testable:
        - Commence par un verbe d'action (Vérifier, Contrôler, S'assurer...)
        - Doit être concret et mesurable
        - Doit couvrir l'intégralité de la règle
        - Formulation claire et concise
        
        Règle: {rule}
        
        Retourne uniquement le PDC sans commentaire.
        """
        
        result = generate_with_azure_openai(prompt, client)
        return result if result else f"Vérifier que {rule}"

def compare_rules_pdc(rules, pdc_list):
    """Compare les règles avec les PDC existants"""
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform(rules + pdc_list)
    similarity = cosine_similarity(tfidf_matrix[:len(rules)], tfidf_matrix[len(rules):])
    return similarity

def create_test_case(pdc, index, is_manual=False, use_ai=False):
    """Crée un cas de test à partir d'un PDC"""
    if not use_ai:
        templates = [
            f"Le système doit satisfaire : {pdc}",
            f"Confirmer que {pdc}",
            f"Tester la conformité de : {pdc}"
        ]
        return {
            "ID": f"CT-{index:03d}",
            "Type": "Manuel" if is_manual else "Auto-généré",
            "PDC": pdc,
            "Description": random.choice(templates) if not is_manual else pdc,
            "Étapes": f"1. Préparer l'environnement\n2. Exécuter: {pdc}\n3. Vérifier le résultat",
            "Résultat attendu": f"{pdc} est correctement implémenté"
        }
    else:
        client = setup_azure_openai()
        if not client:
            return {
                "ID": f"CT-{index:03d}",
                "Type": "Manuel" if is_manual else "Auto-généré",
                "PDC": pdc,
                "Description": pdc,
                "Étapes": f"1. Préparer l'environnement\n2. Exécuter: {pdc}\n3. Vérifier le résultat",
                "Résultat attendu": f"{pdc} est correctement implémenté"
            }
            
        prompt = f"""
        Crée un cas de test complet et détaillé à partir de ce PDC:
        {pdc}
        
        Format de sortie:
        - ID: CT-XXX
        - Type: Manuel/Auto-généré
        - Description: description claire de l'objectif du test
        - Étapes: liste numérotée des actions à réaliser
        - Résultat attendu: résultat observable et vérifiable
        
        Le cas de test doit être réaliste, précis et couvrir tous les aspects du PDC.
        """
        
        result = generate_with_azure_openai(prompt, client)
        if result:
            try:
                # Essaye d'extraire les parties du résultat
                desc = re.search(r"Description:\s*(.+?)\n", result, re.IGNORECASE)
                steps = re.search(r"Étapes:\s*((?:\d\..+?\n)+)", result, re.IGNORECASE)
                expected = re.search(r"Résultat attendu:\s*(.+)", result, re.IGNORECASE)
                
                return {
                    "ID": f"CT-{index:03d}",
                    "Type": "Manuel" if is_manual else "Auto-généré",
                    "PDC": pdc,
                    "Description": desc.group(1).strip() if desc else pdc,
                    "Étapes": steps.group(1).strip() if steps else f"1. Préparer l'environnement\n2. Exécuter: {pdc}\n3. Vérifier le résultat",
                    "Résultat attendu": expected.group(1).strip() if expected else f"{pdc} est correctement implémenté"
                }
            except:
                pass
        
        # Fallback si échec
        return {
            "ID": f"CT-{index:03d}",
            "Type": "Manuel" if is_manual else "Auto-généré",
            "PDC": pdc,
            "Description": pdc,
            "Étapes": f"1. Préparer l'environnement\n2. Exécuter: {pdc}\n3. Vérifier le résultat",
            "Résultat attendu": f"{pdc} est correctement implémenté"
        }

def create_pdc_document(pdc_list):
    """Crée un document Word à partir des PDC"""
    doc = Document()
    doc.add_heading('Points de Contrôle (PDC)', level=1)
    for i, pdc in enumerate(pdc_list, 1):
        p = doc.add_paragraph(style='ListBullet')
        p.add_run(f"{i}. {pdc}").bold = True
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ----------------------------
# INTERFACE UTILISATEUR
# ----------------------------
st.title("Générateur de Cas de Test avec Azure OpenAI")
tab1, tab2, tab3, tab4, tab5 = st.tabs(["📤 Extraction", "🔍 Analyse", "☁️ WordCloud", "📜 Règles", "✅ PDC & Tests"])

with tab1:
    st.header("Extraction de Texte")
    uploaded_file = st.file_uploader("Téléversez un document (PDF ou DOCX)", type=["pdf", "docx"])
    
    if uploaded_file and st.button("Extraire le texte"):
        with st.spinner("Extraction en cours..."):
            extracted_text = extract_text(uploaded_file)
            
            if extracted_text:
                st.session_state.text = extracted_text
                st.success("Texte extrait avec succès !")
                
                with st.expander("Aperçu du texte"):
                    st.text(extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else ""))

with tab2:
    if 'text' not in st.session_state:
        st.warning("Veuillez d'abord extraire un texte dans l'onglet 'Extraction'")
    else:
        nlp_model = load_nlp_model()
        if not nlp_model:
            st.error("Modèle NLP non disponible pour le nettoyage")
        else:
            with st.spinner("Nettoyage approfondi en cours..."):
                st.session_state.text_clean = clean_text(st.session_state.text, nlp_model)
                st.session_state.freq = calculate_frequencies(st.session_state.text_clean)
            
            st.subheader("Fréquence des mots (nettoyés)")
            top_n = st.slider("Nombre de mots à afficher", 5, 50, 20)
            st.dataframe(st.session_state.freq.head(top_n))

with tab3:
    st.header("Visualisation WordCloud")
    
    if 'text_clean' not in st.session_state:
        st.warning("Veuillez d'abord analyser un texte dans l'onglet 'Analyse'")
    else:
        with st.expander("Paramètres avancés"):
            col1, col2 = st.columns(2)
            with col1:
                width = st.slider("Largeur", 400, 1200, 800, key="wc_width")
                height = st.slider("Hauteur", 200, 800, 400, key="wc_height")
            with col2:
                bg_color = st.color_picker("Couleur de fond", "#FFFFFF", key="wc_bg")
                colormap = st.selectbox("Palette", ["viridis", "plasma", "inferno", "magma", "cividis"], key="wc_cmap")
        
        if st.button("Générer le WordCloud"):
            freq_dict = st.session_state.freq.to_dict()
            fig = generate_wordcloud(
                freq_dict,
                width=width,
                height=height,
                background_color=bg_color,
                colormap=colormap
            )
            
            st.pyplot(fig)
            
            img_buffer = BytesIO()
            plt.savefig(img_buffer, format='png', bbox_inches='tight')
            st.download_button(
                label="💾 Télécharger l'image",
                data=img_buffer.getvalue(),
                file_name="wordcloud.png",
                mime="image/png"
            )

with tab4:
    st.header("Extraction des Règles de Gestion")
    nlp_model = load_nlp_model()
    
    # Option pour utiliser Azure OpenAI
    use_ai_rules = st.checkbox("Utiliser Azure OpenAI pour améliorer l'extraction", value=False)
    
    if 'text' not in st.session_state:
        st.warning("Veuillez d'abord extraire un texte dans l'onglet 'Extraction'")
    elif not nlp_model and not use_ai_rules:
        st.error("Le traitement NLP n'est pas disponible")
    else:
        if st.button("Extraire les règles", type="primary"):
            with st.spinner("Analyse en cours (cela peut prendre quelques minutes)..."):
                rules = extract_business_rules(st.session_state.text, nlp_model, use_ai=use_ai_rules)
                
                if rules:
                    st.session_state.rules = rules
                    st.success(f"{len(rules)} règles identifiées !")
                    
                    st.subheader("Règles extraites")
                    items_per_page = 5
                    total_pages = (len(rules) + items_per_page - 1) // items_per_page
                    
                    page = st.number_input("Page", 1, total_pages, 1, 
                                         help="Naviguez entre les pages de résultats")
                    
                    start_idx = (page - 1) * items_per_page
                    end_idx = min(start_idx + items_per_page, len(rules))
                    
                    for i in range(start_idx, end_idx):
                        st.markdown(f"**Règle {i+1}**")
                        st.info(rules[i])
                    
                    st.subheader("Export des résultats")
                    docx_file = create_rules_document(rules)
                    st.download_button(
                        "📄 Télécharger au format Word",
                        data=docx_file,
                        file_name="regles_gestion.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    with st.expander("Analyse avancée"):
                        st.metric("Nombre total de règles", len(rules))
                        avg_length = sum(len(rule.split()) for rule in rules) / len(rules)
                        st.metric("Longueur moyenne des règles", f"{avg_length:.1f} mots")
                else:
                    st.warning("Aucune règle de gestion n'a été identifiée dans le document")

with tab5:
    st.header("Gestion des Points de Contrôle et Cas de Test")
    
    # Options Azure OpenAI
    col1, col2 = st.columns(2)
    with col1:
        use_ai_pdc = st.checkbox("Utiliser Azure OpenAI pour les PDC", value=False)
    with col2:
        use_ai_tests = st.checkbox("Utiliser Azure OpenAI pour les cas de test", value=False)
    
    if 'rules' not in st.session_state:
        st.warning("Veuillez d'abord extraire les règles dans l'onglet 'Règles'")
    else:
        # Section 1: Chargement des PDC existants
        st.subheader("1. Chargement des PDC existants")
        has_pdc = st.radio("Avez-vous des PDC existants à importer ?", 
                          ("Oui, j'ai des PDC existants", "Non, générer des PDC automatiquement"),
                          index=0)
        
        pdc_file = None
        pdc_text = ""
        
        if has_pdc.startswith("Oui"):
            pdc_file = st.file_uploader("Téléversez votre fichier PDC (PDF/DOCX/TXT)", 
                                       type=["pdf", "docx", "txt"], 
                                       key="pdc_uploader")
            
            if pdc_file:
                with st.spinner("Extraction des PDC en cours..."):
                    pdc_text = extract_text(pdc_file)
                    st.session_state.pdc_list = extract_pdc_from_text(pdc_text)
                    
                    if st.session_state.pdc_list:
                        st.success(f"{len(st.session_state.pdc_list)} PDC extraits !")
                        with st.expander("Aperçu des PDC"):
                            for i, pdc in enumerate(st.session_state.pdc_list[:5], 1):
                                st.markdown(f"{i}. {pdc}")
                    else:
                        st.warning("Aucun PDC détecté dans le document")
                        st.session_state.pdc_list = []
        
        # Section 2: Génération des PDC
        st.subheader("2. Génération des PDC")
        if has_pdc.startswith("Non") or (has_pdc.startswith("Oui") and pdc_file):
            if st.button("Générer/Compléter les PDC", type="primary"):
                with st.spinner("Création des PDC..."):
                    # Initialisation de la liste PDC
                    if 'pdc_list' not in st.session_state:
                        st.session_state.pdc_list = []
                    
                    # Pour les règles sans PDC correspondant
                    if has_pdc.startswith("Oui") and pdc_file:
                        similarity = compare_rules_pdc(st.session_state.rules, st.session_state.pdc_list)
                        threshold = st.slider("Seuil de similarité pour les correspondances", 0.1, 1.0, 0.6)
                        
                        for i, rule in enumerate(st.session_state.rules):
                            if similarity[i].max() < threshold:
                                generated_pdc = generate_pdc_from_rule(rule, use_ai=use_ai_pdc)
                                st.session_state.pdc_list.append(generated_pdc)
                    else:
                        # Génération automatique complète
                        st.session_state.pdc_list = [generate_pdc_from_rule(rule, use_ai=use_ai_pdc) for rule in st.session_state.rules]
                    
                    st.success(f"{len(st.session_state.pdc_list)} PDC prêts !")
        
        # Section 3: Visualisation et Export
        if 'pdc_list' in st.session_state and st.session_state.pdc_list:
            st.subheader("3. Points de Contrôle")
            
            # Affichage paginé
            pdc_per_page = 5
            total_pages = (len(st.session_state.pdc_list) + pdc_per_page - 1) // pdc_per_page
            page = st.number_input("Page", 1, total_pages, 1)
            
            start_idx = (page - 1) * pdc_per_page
            end_idx = min(start_idx + pdc_per_page, len(st.session_state.pdc_list))
            
            for i in range(start_idx, end_idx):
                st.markdown(f"**PDC {i+1}**")
                st.info(st.session_state.pdc_list[i])
            
            # Export PDC
            st.download_button(
                "📥 Télécharger les PDC (DOCX)",
                data=create_pdc_document(st.session_state.pdc_list),
                file_name="points_de_controle.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Section 4: Génération des Cas de Test
            st.subheader("4. Cas de Test Associés")
            
            if st.button("Générer les Cas de Test"):
                with st.spinner("Création des cas de test..."):
                    st.session_state.test_cases = []
                    
                    for i, pdc in enumerate(st.session_state.pdc_list, 1):
                        is_manual = has_pdc.startswith("Oui") and i <= len(st.session_state.pdc_list)
                        st.session_state.test_cases.append(create_test_case(pdc, i, is_manual, use_ai=use_ai_tests))
                    
                    st.success(f"{len(st.session_state.test_cases)} cas de test générés !")
            
            # Affichage des Cas de Test
            if 'test_cases' in st.session_state:
                df_test_cases = pd.DataFrame(st.session_state.test_cases)
                st.dataframe(df_test_cases[["ID", "Type", "PDC", "Description"]])
                
                # Export des Cas de Test
                test_cases_doc = Document()
                test_cases_doc.add_heading('Cas de Test', level=1)
                
                table = test_cases_doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                headers = ["ID", "Type", "PDC", "Description", "Étapes"]
                for i, header in enumerate(headers):
                    table.cell(0, i).text = header
                
                for case in st.session_state.test_cases:
                    row = table.add_row().cells
                    row[0].text = case["ID"]
                    row[1].text = case["Type"]
                    row[2].text = case["PDC"]
                    row[3].text = case["Description"]
                    row[4].text = case["Étapes"]
                
                buffer = BytesIO()
                test_cases_doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    "📥 Télécharger les Cas de Test (DOCX)",
                    data=buffer,
                    file_name="cas_de_test.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ----------------------------
# PIED DE PAGE
# ----------------------------
st.markdown("---")
st.caption("Application développée avec Streamlit - Mise à jour : %s" % pd.Timestamp.now().strftime("%d/%m/%Y"))
